from django.contrib.auth.decorators import login_required
from django.http import HttpResponse, HttpRequest, Http404
from django.shortcuts import render, redirect
from django.views.generic import ListView
from Student_module.models import Student, Course, Semester
from io import BytesIO
from django.views import View
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
from django.views.decorators.csrf import csrf_exempt
import json
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment
from docx import Document
from django.utils.decorators import method_decorator


@login_required
def Home_page(request):
    return redirect('home')


@method_decorator(login_required, name="dispatch")
class HomeView(ListView):
    template_name = "home-page.html"
    model = Student
    paginate_by = 1

    def get_context_data(self):
        context = super().get_context_data()
        user = self.request.user.username
        semester_number = self.kwargs.get("sn")
        if semester_number is not None:
            context["courses"] = Course.objects.filter(
                semester__student__username__iexact=user,
                semester__semester_number__iexact=semester_number,
            ).order_by("-grade")
        else:
            context["courses"] = Course.objects.filter(
                semester__student__username__iexact=user
            ).order_by("-grade")
        context["student"] = Student.objects.filter(username__iexact=user)
        context["semesters"] = Semester.objects.filter(student__username__iexact=user)
        return context


@login_required
def FilterData(request: HttpRequest):
    semester_number = int(request.GET.get("semester"))
    grade = int(request.GET.get("grade"))
    user = request.user.username
    student = Student.objects.filter(username__iexact=user)
    semesters = Semester.objects.filter(student__username__iexact=user)
    if grade and semester_number is not None:
        if grade == 17:
            course = Course.objects.filter(
                semester__semester_number__iexact=semester_number, grade__gte=grade,
                semester__student__username__iexact=user
            ).order_by('-grade')
        else:
            course = Course.objects.filter(
                semester__semester_number__iexact=semester_number, grade__lt=grade,
                semester__student__username__iexact=user
            ).order_by('-grade')

        context = {
            'courses': course,
            'student': student,
            'semesters': semesters
        }
        return render(request, 'home-partial.html', context)


@method_decorator(csrf_exempt, name='dispatch')
class DownloadPDFView(View):
    def post(self, request):
        data = json.loads(request.body)
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="report.pdf"'
        print(data)
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        elements = []
        try:
            boolean = bool(data[0]['gpa'])
            table_data = [['Semester Number', 'Gpa', 'Total Credits', 'Passed Credits', 'Failed Credits']]
            for item in data:
                table_data.append(
                    [item['semester_number'], item['gpa'], item['total_credits'], item['passed_credits'],
                     item['failed_credits']])

            table = Table(table_data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ]))

            elements.append(table)
            doc.build(elements)

            pdf = buffer.getvalue()
            buffer.close()
            response.write(pdf)
            return response

        except:
            table_data = [['Course Name', 'Grade', 'Rank', 'Credit', 'Professor']]
            for item in data:
                table_data.append([item['name'], item['grade'], item['rank'], item['credit'], item['professor']])

            table = Table(table_data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ]))

            elements.append(table)
            doc.build(elements)

            pdf = buffer.getvalue()
            buffer.close()
            response.write(pdf)
            return response


@method_decorator(csrf_exempt, name='dispatch')
class DownloadWordView(View):
    def post(self, request):
        data = json.loads(request.body)
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="report.docx"'

        doc = Document()

        try:
            boolean = bool(data[0]['gpa'])
            doc.add_heading('Semester Report', 0)
            table = doc.add_table(rows=1, cols=5)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Semester Number'
            hdr_cells[1].text = 'GPA'
            hdr_cells[2].text = 'Total Credits'
            hdr_cells[3].text = 'Passed Credits'
            hdr_cells[4].text = 'Failed Credits'
            for item in data:
                row_cells = table.add_row().cells
                row_cells[0].text = str(item['semester_number'])
                row_cells[1].text = str(item['gpa'])
                row_cells[2].text = str(item['total_credits'])
                row_cells[3].text = str(item['passed_credits'])
                row_cells[4].text = str(item['failed_credits'])

        except:
            doc.add_heading('Course Report', 0)
            table = doc.add_table(rows=1, cols=5)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Course Name'
            hdr_cells[1].text = 'Grade'
            hdr_cells[2].text = 'Rank'
            hdr_cells[3].text = 'Credit'
            hdr_cells[4].text = 'Professor'
            for item in data:
                row_cells = table.add_row().cells
                row_cells[0].text = item['name']
                row_cells[1].text = str(item['grade'])
                row_cells[2].text = str(item['rank'])
                row_cells[3].text = str(item['credit'])
                row_cells[4].text = item['professor']

        doc.save(response)
        return response


@method_decorator(csrf_exempt, name='dispatch')
class DownloadExcelView(View):
    def post(self, request):
        data = json.loads(request.body)
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = 'attachment; filename="report.xlsx"'
        try:
            boolean = bool(data[0]['gpa'])
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = 'Courses'

            headers = ['Semester Number', 'Gpa', 'Total Credits', 'Passed Credits', 'Failed Credits']
            sheet.append(headers)
            for col_num, header in enumerate(headers, 1):
                cell = sheet.cell(row=1, column=col_num)
                cell.font = Font(bold=True)
                cell.alignment = Alignment(horizontal='center')

            for item in data:
                sheet.append([item['semester_number'], item['gpa'], item['total_credits'], item['passed_credits'],
                              item['failed_credits']])

            for row in sheet.iter_rows(min_row=2, max_col=5, max_row=sheet.max_row):
                for cell in row:
                    cell.alignment = Alignment(horizontal='center')

            workbook.save(response)
            return response
        except:
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = 'Courses'

            headers = ['Course Name', 'Grade', 'Rank', 'Credit', 'Professor']
            sheet.append(headers)

            for col_num, header in enumerate(headers, 1):
                cell = sheet.cell(row=1, column=col_num)
                cell.font = Font(bold=True)
                cell.alignment = Alignment(horizontal='center')

            for item in data:
                sheet.append([item['name'], item['grade'], item['rank'], item['credit'], item['professor']])

            for row in sheet.iter_rows(min_row=2, max_col=5, max_row=sheet.max_row):
                for cell in row:
                    cell.alignment = Alignment(horizontal='center')

            workbook.save(response)
            return response


@login_required
def Overall(request: HttpRequest):
    user = request.user.username
    student = Student.objects.filter(username__iexact=user)
    semester = Semester.objects.filter(student__username__iexact=user).order_by('-gpa')
    context = {
        'student': student,
        'semesters': semester
    }
    return render(request, 'overall-page.html', context)


